"""Smart Resume Screening and Candidate Ranking Tool.

A lightweight FastAPI application that extracts resume text, identifies skills,
compares candidates with a job description, and returns an explainable ranking.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import HTMLResponse
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

app = FastAPI(
    title="Smart Resume Screening and Candidate Ranking Tool",
    version="1.0.0",
    description="NLP-based resume screening and candidate ranking API.",
)

# A practical starter vocabulary. The matcher is intentionally transparent so
# the project can be extended with a trained NER/embedding model later.
SKILLS = {
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "go", "rust",
    "sql", "postgresql", "mysql", "mongodb", "redis", "html", "css", "react",
    "next.js", "node.js", "express", "fastapi", "flask", "django", "spring",
    "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch",
    "keras", "opencv", "spacy", "transformers", "sentence transformers",
    "git", "github", "docker", "kubernetes", "aws", "azure", "gcp", "linux",
    "rest api", "graphql", "microservices", "power bi", "tableau", "excel",
    "data analysis", "data science", "statistics", "communication", "leadership",
    "problem solving", "teamwork", "agile", "scrum",
}

STOP_WORDS = {
    "and", "the", "with", "for", "from", "that", "this", "have", "has", "are",
    "you", "your", "our", "will", "job", "role", "work", "years", "year", "using",
    "into", "about", "their", "they", "them", "who", "what", "where", "when",
}


def normalize(text: str) -> str:
    """Normalize document text before feature extraction."""
    text = text.lower().replace("\u2013", "-").replace("\u2014", "-")
    return re.sub(r"\s+", " ", text).strip()


def extract_pdf(data: bytes) -> str:
    """Extract text from a PDF in memory without writing candidate data to disk."""
    document = fitz.open(stream=data, filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in document)
    finally:
        document.close()


def extract_docx(data: bytes) -> str:
    """Extract paragraphs and table text from a DOCX document."""
    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch document extraction based on the uploaded extension."""
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return extract_pdf(data)
    if extension == ".docx":
        return extract_docx(data)
    if extension == ".txt":
        return data.decode("utf-8", errors="ignore")
    raise HTTPException(status_code=400, detail=f"Unsupported file type: {extension}")


def extract_skills(text: str) -> list[str]:
    """Return skills found in text using phrase-aware matching."""
    normalized = normalize(text)
    found = []
    for skill in sorted(SKILLS, key=len, reverse=True):
        pattern = rf"(?<![a-z0-9+#.-]){re.escape(skill.lower())}(?![a-z0-9+#.-])"
        if re.search(pattern, normalized):
            found.append(skill)
    return sorted(set(found))


def extract_email(text: str) -> str | None:
    match = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.I)
    return match.group(0) if match else None


def extract_phone(text: str) -> str | None:
    match = re.search(r"(?:\+?\d[\d\s().-]{8,}\d)", text)
    return re.sub(r"\s+", " ", match.group(0)).strip() if match else None


def extract_experience_years(text: str) -> float | None:
    patterns = [
        r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s*(?:of)?\s*experience",
        r"experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*\+?\s*years?",
    ]
    normalized = normalize(text)
    for pattern in patterns:
        match = re.search(pattern, normalized)
        if match:
            return float(match.group(1))
    return None


def tokenize(text: str) -> list[str]:
    """Create simple tokens for lightweight keyword analysis."""
    return [token for token in re.findall(r"[a-z][a-z0-9+#.-]{1,}", normalize(text)) if token not in STOP_WORDS]


def score_candidate(job_description: str, resume_text: str) -> dict[str, Any]:
    """Calculate an explainable hybrid score from semantic and skill signals."""
    job = normalize(job_description)
    resume = normalize(resume_text)

    vectorizer = TfidfVectorizer(ngram_range=(1, 2), min_df=1, max_features=5000)
    matrix = vectorizer.fit_transform([job, resume])
    semantic_score = float(cosine_similarity(matrix[0:1], matrix[1:2])[0][0])

    job_skills = set(extract_skills(job))
    resume_skills = set(extract_skills(resume))
    matched = sorted(job_skills & resume_skills)
    missing = sorted(job_skills - resume_skills)
    skill_score = len(matched) / len(job_skills) if job_skills else semantic_score

    job_tokens = set(tokenize(job))
    resume_tokens = set(tokenize(resume))
    keyword_score = len(job_tokens & resume_tokens) / len(job_tokens) if job_tokens else 0.0

    # Transparent weighting makes the result easy to explain to a recruiter.
    overall = (semantic_score * 0.45) + (skill_score * 0.40) + (keyword_score * 0.15)
    overall = round(max(0.0, min(1.0, overall)) * 100, 1)

    return {
        "score": overall,
        "semantic_similarity": round(semantic_score * 100, 1),
        "skill_match": round(skill_score * 100, 1),
        "keyword_match": round(keyword_score * 100, 1),
        "matched_skills": matched,
        "missing_skills": missing,
        "resume_skills": sorted(resume_skills),
        "experience_years": extract_experience_years(resume_text),
    }


@app.get("/health")
def health() -> dict[str, str]:
    """Health endpoint for local checks and deployment platforms."""
    return {"status": "ok", "service": "smart-resume-screening"}


@app.post("/api/analyze")
async def analyze_resumes(
    job_description: str = Form(...),
    resumes: list[UploadFile] = File(...),
) -> dict[str, Any]:
    """Analyze multiple resumes against one job description and rank them."""
    if len(job_description.strip()) < 30:
        raise HTTPException(status_code=400, detail="Job description must contain at least 30 characters.")
    if not resumes:
        raise HTTPException(status_code=400, detail="Upload at least one resume.")
    if len(resumes) > 25:
        raise HTTPException(status_code=400, detail="Maximum 25 resumes per analysis.")

    results = []
    for upload in resumes:
        data = await upload.read()
        if len(data) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail=f"{upload.filename} exceeds the 5 MB limit.")
        text = extract_text(upload.filename or "resume.txt", data)
        if len(text.strip()) < 50:
            results.append({
                "filename": upload.filename,
                "score": 0.0,
                "error": "Not enough readable text was found in this document.",
            })
            continue

        scored = score_candidate(job_description, text)
        scored.update({
            "filename": upload.filename,
            "email": extract_email(text),
            "phone": extract_phone(text),
            "preview": re.sub(r"\s+", " ", text).strip()[:240],
        })
        results.append(scored)

    results.sort(key=lambda item: item.get("score", 0), reverse=True)
    for index, result in enumerate(results, start=1):
        result["rank"] = index

    return {
        "job_title": next((line.strip() for line in job_description.splitlines() if line.strip()), "Job Position")[:100],
        "candidate_count": len(results),
        "results": results,
    }


INDEX_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8" />
<meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>Smart Resume Screening</title>
<style>
:root{font-family:Inter,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;color:#172033;background:#f5f7fb}
*{box-sizing:border-box}body{margin:0}.shell{max-width:1180px;margin:0 auto;padding:42px 22px}.header{margin-bottom:28px}.eyebrow{font-size:13px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:#5b6b86}.header h1{margin:7px 0;font-size:34px}.header p{margin:0;color:#68758b;max-width:720px;line-height:1.6}.grid{display:grid;grid-template-columns:390px 1fr;gap:22px}.card{background:#fff;border:1px solid #e3e8f0;border-radius:14px;padding:22px;box-shadow:0 5px 18px rgba(31,45,61,.05)}label{display:block;font-size:13px;font-weight:700;margin:0 0 8px;color:#334155}textarea{width:100%;min-height:280px;border:1px solid #d8dee8;border-radius:10px;padding:13px;resize:vertical;font:inherit;outline:none}textarea:focus{border-color:#5b6b86}.drop{border:1.5px dashed #b8c2d2;border-radius:10px;padding:18px;text-align:center;margin-top:16px;background:#fafbfd}.drop input{width:100%}.button{width:100%;margin-top:16px;border:0;border-radius:9px;padding:12px 16px;background:#172033;color:#fff;font-weight:700;cursor:pointer}.button:disabled{opacity:.55;cursor:wait}.status{margin-top:10px;color:#68758b;font-size:13px}.empty{height:100%;min-height:430px;display:flex;align-items:center;justify-content:center;color:#7b8799;text-align:center}.results{display:none}.summary{display:flex;gap:12px;margin-bottom:18px}.stat{flex:1;border:1px solid #e3e8f0;border-radius:10px;padding:14px;background:#fafbfd}.stat b{display:block;font-size:22px;color:#172033}.stat span{font-size:12px;color:#68758b}.candidate{border:1px solid #e3e8f0;border-radius:12px;padding:17px;margin-bottom:12px}.candidate-top{display:flex;justify-content:space-between;gap:12px}.candidate h3{margin:0 0 4px;font-size:16px}.muted{font-size:12px;color:#718096}.score{font-size:22px;font-weight:800}.bar{height:7px;background:#e9edf3;border-radius:10px;margin:12px 0;overflow:hidden}.fill{height:100%;background:#172033;border-radius:10px}.chips{display:flex;flex-wrap:wrap;gap:6px}.chip{font-size:11px;padding:5px 8px;border-radius:999px;background:#eef2f7;color:#344054}.chip.missing{background:#fff0f0;color:#a23b3b}.detail{margin-top:10px;font-size:12px;color:#5f6c80}.error{color:#a23b3b}.api{margin-top:24px;font-size:12px;color:#8490a3}@media(max-width:850px){.grid{grid-template-columns:1fr}.summary{flex-wrap:wrap}.stat{min-width:130px}}
</style></head>
<body><main class="shell">
<header class="header"><div class="eyebrow">Recruitment Intelligence</div><h1>Smart Resume Screening</h1><p>Compare resumes against a job description, identify relevant skills, and rank candidates using a transparent NLP-based scoring model.</p></header>
<section class="grid">
<div class="card"><form id="form"><label for="job">Job Description</label><textarea id="job" name="job_description" placeholder="Paste the complete job description here..."></textarea><div class="drop"><label for="resumes">Resume files</label><input id="resumes" name="resumes" type="file" accept=".pdf,.docx,.txt" multiple /></div><button class="button" id="submit" type="submit">Screen Resumes</button><div class="status" id="status"></div></form><div class="api">Supported: PDF, DOCX, TXT · Maximum 25 files · 5 MB per file</div></div>
<div class="card"><div class="empty" id="empty">Upload a job description and resumes to see the ranked candidate list.</div><div class="results" id="results"><div class="summary"><div class="stat"><b id="count">0</b><span>Candidates</span></div><div class="stat"><b id="top">0%</b><span>Top match</span></div><div class="stat"><b id="skills">0</b><span>Job skills found</span></div></div><div id="list"></div></div></div>
</section></main>
<script>
const form=document.getElementById('form'), submit=document.getElementById('submit'), status=document.getElementById('status');
form.addEventListener('submit',async e=>{e.preventDefault();const files=document.getElementById('resumes').files;const job=document.getElementById('job').value.trim();if(!job||job.length<30){status.textContent='Enter a job description of at least 30 characters.';return}if(!files.length){status.textContent='Select at least one resume.';return}const data=new FormData();data.append('job_description',job);[...files].forEach(f=>data.append('resumes',f));submit.disabled=true;status.textContent='Analyzing resumes...';try{const res=await fetch('/api/analyze',{method:'POST',body:data});const body=await res.json();if(!res.ok)throw new Error(body.detail||'Analysis failed');render(body);status.textContent=`Analysis complete · ${body.candidate_count} candidate(s)`}catch(err){status.textContent=err.message;status.className='status error'}finally{submit.disabled=false}});
function render(data){document.getElementById('empty').style.display='none';document.getElementById('results').style.display='block';document.getElementById('count').textContent=data.candidate_count;document.getElementById('top').textContent=(data.results[0]?.score||0)+'%';const jobSkills=new Set(data.results.flatMap(x=>x.matched_skills||[]).concat(data.results.flatMap(x=>x.missing_skills||[])));document.getElementById('skills').textContent=jobSkills.size;document.getElementById('list').innerHTML=data.results.map(c=>{if(c.error)return `<article class="candidate"><div class="candidate-top"><div><h3>#${c.rank} ${escapeHtml(c.filename)}</h3><div class="muted">${escapeHtml(c.error)}</div></div></div></article>`;return `<article class="candidate"><div class="candidate-top"><div><h3>#${c.rank} ${escapeHtml(c.filename)}</h3><div class="muted">${escapeHtml(c.email||'Email not detected')} · ${escapeHtml(c.experience_years!=null?c.experience_years+' years experience':'Experience not detected')}</div></div><div class="score">${c.score}%</div></div><div class="bar"><div class="fill" style="width:${c.score}%"></div></div><div class="detail">Semantic ${c.semantic_similarity}% · Skills ${c.skill_match}% · Keywords ${c.keyword_match}%</div><p class="detail"><b>Matched skills</b></p><div class="chips">${(c.matched_skills||[]).map(s=>`<span class="chip">${escapeHtml(s)}</span>`).join('')||'<span class="muted">No explicit matches</span>'}</div><p class="detail"><b>Skill gaps</b></p><div class="chips">${(c.missing_skills||[]).map(s=>`<span class="chip missing">${escapeHtml(s)}</span>`).join('')||'<span class="muted">None detected</span>'}</div></article>`}).join('')}
function escapeHtml(value){return String(value??'').replace(/[&<>'"]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;',"'":'&#39;','"':'&quot;'}[c]))}
</script></body></html>"""


@app.get("/", response_class=HTMLResponse)
def index() -> str:
    """Serve the self-contained recruiter interface."""
    return INDEX_HTML
